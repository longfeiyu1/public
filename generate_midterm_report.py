from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


COURSES = [
    ("学术英语写作", "公共必修课", 92, 2),
    ("大语言模型驱动多智能体", "专业选修课", 89, 2),
    ("新时代中国特色社会主义理论与实践", "公共必修课", 83, 2),
    ("Python程序设计", "公共选修课", 87, 2),
    ("高级机器学习", "专业必修课", 84, 3),
    ("高级数据库技术", "学科基础课", 75, 3),
    ("数理逻辑基础", "学科基础课", 61, 3),
    ("信息检索与互联网挖掘", "专业选修课", 90, 2),
    ("人工智能概述", "公共选修课", 90, 2),
    ("人工智能研究方法与论文写作", "专业选修课", 84, 2),
    ("高级数据挖掘", "专业必修课", 87, 3),
    ("计算机视觉", "专业选修课", 89, 2),
    ("最优化方法", "学科基础课", 61, 3),
    ("习近平法治思想专题研究", "公选思政课", 83, 1),
    ("美育教育课程", "公共选修课", 90, 2),
    ("高级算法分析与设计", "专业必修课", 88, 3),
]


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_format(paragraph, first_line_chars=2):
    paragraph.paragraph_format.first_line_indent = Cm(0.74 * first_line_chars)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)


def set_document_base_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(16)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(12)


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    set_paragraph_format(p)
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)


def add_info_table(doc):
    table = doc.add_table(rows=4, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    widths = [Cm(2.8), Cm(5.2), Cm(2.8), Cm(5.2)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]

    fields = [
        ("学校名称", "【请填写】", "学院", "【请填写】"),
        ("姓名", "【请填写】", "学号", "【请填写】"),
        ("专业", "计算机科学与技术", "导师姓名", "【请填写】"),
        ("考核学年/时间", "【请填写】", "研究方向", "图像修复"),
    ]

    for row, values in zip(table.rows, fields):
        for i, value in enumerate(values):
            set_cell_text(row.cells[i], value, bold=(i % 2 == 0))


def add_course_table(doc):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("课程学习与成绩汇总")
    run.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(12)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["课程名称", "课程性质", "成绩", "学分"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True)

    for course_name, course_type, score, credit in COURSES:
        row = table.add_row().cells
        set_cell_text(row[0], course_name, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row[1], course_type)
        set_cell_text(row[2], str(score))
        set_cell_text(row[3], str(credit))

    summary = table.add_row().cells
    set_cell_text(summary[0], "合计/说明", bold=True)
    set_cell_text(summary[1], "已完成全部课程学习")
    set_cell_text(summary[2], "平均分约83.31")
    set_cell_text(summary[3], "总学分37")


def set_page_margins(section):
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)


def add_page_number(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def build_document():
    doc = Document()
    set_document_base_style(doc)
    section = doc.sections[0]
    set_page_margins(section)
    add_page_number(section)

    add_title(doc, "硕士研究生中期考核报告")
    doc.add_paragraph()
    add_info_table(doc)

    add_section_heading(doc, "一、思想品德表现")
    add_body_paragraph(
        doc,
        "本人能够坚持正确的政治方向，认真学习党的路线方针政策，遵守国家法律法规和学校各项规章制度，思想态度端正，集体观念较强。在研究生阶段，能够按照学校要求参加相关思想政治教育活动，注重将理论学习与日常学习生活相结合。在学习和工作中能够保持较为严谨、踏实的态度，尊重师长，团结同学，具有较强的纪律意识和责任意识。"
    )
    add_body_paragraph(
        doc,
        "在学术规范方面，本人能够认识到科研诚信和学术道德的重要性，重视文献阅读、资料整理和研究过程中的规范表达，做到诚实守信、实事求是，未出现违反学术纪律和校纪校规的情况。"
    )

    add_section_heading(doc, "二、课程学习情况")
    add_body_paragraph(
        doc,
        "根据培养计划要求，本人已完成培养方案规定的全部课程学习，课程总学分为37学分，课程结构涵盖公共必修课、公共选修课、学科基础课、专业必修课、专业选修课及思政类选修课程，整体达到培养方案要求，课程成绩均合格。已修课程平均成绩约为83.31分。"
    )
    add_body_paragraph(
        doc,
        "从课程内容看，已完成的课程能够较好支撑计算机科学与技术专业硕士阶段的知识体系构建。其中，高级机器学习、高级数据挖掘、计算机视觉、高级算法分析与设计等课程，强化了本人在人工智能、模式识别、算法设计与分析等方面的基础；高级数据库技术、数理逻辑基础、最优化方法等课程，则进一步夯实了专业基础能力。相关课程学习为后续围绕图像修复方向开展研究提供了必要的理论准备和方法支撑。"
    )
    add_body_paragraph(
        doc,
        "在学习过程中，本人能够按时完成课程学习任务，较为重视基础理论与实际应用的结合。对于部分难度较高的基础课程，虽然成绩表现较为一般，但均已达到合格要求，也反映出本人对相关基础知识仍需继续巩固和深化。总体来看，课程学习完成情况符合中期考核要求。"
    )

    add_section_heading(doc, "三、科研学习与课题进展")
    add_body_paragraph(
        doc,
        "目前本人尚无公开发表的科研成果，也未形成阶段性论文或专利成果。结合当前实际情况，本人在研究生阶段的主要工作仍以专业基础积累、研究方向聚焦和方法学习为主。研究方向拟围绕图像修复相关问题展开，重点关注图像退化建模、图像内容恢复以及深度学习方法在图像修复任务中的应用。"
    )
    add_body_paragraph(
        doc,
        "现阶段，本人已结合课程学习对机器学习、计算机视觉和数据处理等内容进行了较系统的梳理，并对图像修复方向的相关研究进行了初步了解和文献阅读。通过前期学习，对该方向的基本问题、常见技术路线和可能面临的难点有了一定认识，但在研究问题凝练、实验方案设计和成果产出方面仍处于起步阶段。"
    )
    add_body_paragraph(
        doc,
        "下一步将进一步在导师指导下明确具体选题，持续开展文献阅读和实验准备工作，逐步推进研究工作从理论学习向问题分析、方法实现和实验验证过渡，争取在后续阶段形成较为清晰的研究思路和阶段性成果。"
    )

    add_section_heading(doc, "四、实践与实习情况")
    add_body_paragraph(
        doc,
        "本人在研究生阶段具有一定实习经历。通过参与相关实践工作，进一步体会到专业理论知识在实际场景中的应用价值，也更加认识到工程实现能力、问题分析能力和团队协作能力对于计算机类研究生培养的重要性。"
    )
    add_body_paragraph(
        doc,
        "实习过程在一定程度上强化了本人对“学”与“用”相结合的理解，使本人能够从实际需求出发审视专业知识的应用方式，增强了将课程学习内容与工程实践相结合的意识。相关经历对后续研究工作的开展也具有积极作用，有助于在课题研究中更加重视问题背景、技术可行性和结果落地性。"
    )

    add_section_heading(doc, "五、存在的不足")
    add_body_paragraph(
        doc,
        "结合前一阶段的学习情况，本人仍存在以下不足：一是科研起步相对较慢，尚未形成明确、稳定的阶段性成果；二是部分基础课程掌握还不够扎实，需要进一步加强对数理基础和优化方法等内容的理解；三是在学术论文阅读、问题提炼、实验设计和研究总结方面，还需要持续训练和积累。"
    )

    add_section_heading(doc, "六、下一阶段学习与研究计划")
    add_body_paragraph(
        doc,
        "下一阶段，本人将继续按照培养计划要求，在导师指导下推进学位论文相关工作。首先，进一步聚焦图像修复方向，明确研究问题和技术路线；其次，加强对国内外相关文献的系统梳理，提升对研究现状和关键问题的把握；再次，结合已有课程基础和实践经历，逐步开展算法实现与实验验证，提升科研训练的规范性和有效性。"
    )
    add_body_paragraph(
        doc,
        "同时，本人将继续巩固专业基础，特别是针对数理基础、优化方法和科研写作等薄弱环节进行有针对性的补强；在实践方面，继续注重理论联系实际，努力提高分析问题、解决问题和总结归纳的能力，争取顺利完成后续培养任务。"
    )

    add_section_heading(doc, "七、自我总结")
    add_body_paragraph(
        doc,
        "总体来看，本人在中期考核前已完成培养计划规定的课程学习任务，课程成绩全部合格，学分和学分结构达到培养方案要求；在思想品德、纪律作风和日常学习方面表现较为稳定；在科研方面虽然尚未取得实质性成果，但已初步确定研究方向，并具备继续深入开展研究的基础。后续本人将坚持实事求是、踏实推进的原则，在导师指导下认真完成研究生阶段各项培养任务。"
    )

    add_body_paragraph(doc, "研究生签名：______________        日期：______________")
    add_body_paragraph(doc, "导师意见：____________________________________________________________")
    add_body_paragraph(doc, "____________________________________________________________")
    add_body_paragraph(doc, "导师签名：______________        日期：______________")

    add_course_table(doc)
    return doc


if __name__ == "__main__":
    document = build_document()
    output_path = "硕士研究生中期考核报告（可填写版）.docx"
    document.save(output_path)
    print(f"generated: {output_path}")
